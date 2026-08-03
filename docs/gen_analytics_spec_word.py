"""Generate Analytics Dashboard business logic spec as Word (.docx)."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

# ── Color palette ─────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1F, 0x38, 0x64)
BLUE   = RGBColor(0x2F, 0x54, 0x96)
LBLUE  = RGBColor(0xD6, 0xE4, 0xF0)
YELLOW = RGBColor(0xFF, 0xF2, 0xCC)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
BLACK  = RGBColor(0x00, 0x00, 0x00)
LGRAY  = RGBColor(0xF2, 0xF2, 0xF2)

def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    hex_color = "{:02X}{:02X}{:02X}".format(rgb[0], rgb[1], rgb[2])
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_text(cell, text, bold=False, color=BLACK, size=10,
                  align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color

def header_row(table, texts, bg=NAVY, color=WHITE, bold=True, size=10):
    row = table.rows[0]
    for i, text in enumerate(texts):
        cell = row.cells[i]
        set_cell_bg(cell, bg)
        set_cell_text(cell, text, bold=bold, color=color, size=size,
                      align=WD_ALIGN_PARAGRAPH.CENTER)

def add_data_row(table, texts, bg=WHITE, bold_first=False):
    row = table.add_row()
    for i, text in enumerate(texts):
        cell = row.cells[i]
        set_cell_bg(cell, bg)
        bold = bold_first and i == 0
        set_cell_text(cell, text, bold=bold)

def set_col_widths(table, widths_cm):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                cell.width = Cm(widths_cm[i])

def add_section_title(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = WHITE
    # background via shading on paragraph — use heading style workaround
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "2F5496")
    pPr.append(shd)
    p.paragraph_format.left_indent = Cm(0.3)

# ══════════════════════════════════════════════════════════════════════════════
# Title
# ══════════════════════════════════════════════════════════════════════════════
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run("Analytics Dashboard — Business Logic Specification")
title_run.bold = True
title_run.font.size = Pt(16)
title_run.font.color.rgb = NAVY

doc.add_paragraph()  # spacer

# ══════════════════════════════════════════════════════════════════════════════
# Section 1 — Filter Controls
# ══════════════════════════════════════════════════════════════════════════════
add_section_title(doc, "1. Filter Controls")

t = doc.add_table(rows=1, cols=2)
t.style = "Table Grid"
header_row(t, ["Filter", "ค่าที่เลือกได้"])
set_col_widths(t, [5, 10.5])

filters = [
    ("Year Filter",     "All Years / ปีการศึกษา เช่น 2025/2026, 2024/2025"),
    ("Term Filter",     "All Terms / Term 1 / Term 2 / Term 3"),
    ("Fee Type Filter", "All Fee Types / Tuition / ECA / Trip / Exam / Bus"),
]
for i, (k, v) in enumerate(filters):
    bg = LBLUE if i % 2 == 0 else WHITE
    add_data_row(t, [k, v], bg=bg, bold_first=True)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# Section 2 — Stat Cards
# ══════════════════════════════════════════════════════════════════════════════
add_section_title(doc, "2. Stat Cards (แถวบนสุด)")

t = doc.add_table(rows=1, cols=2)
t.style = "Table Grid"
header_row(t, ["Stat Card", "คำอธิบาย"])
set_col_widths(t, [4, 11.5])

stats = [
    ("Gross Revenue",   "รายได้รวมก่อนหักส่วนลด"),
    ("Net Revenue",     "รายได้สุทธิหลังหักส่วนลดและ Bank Fees"),
    ("Bank Fees",       "ค่าธรรมเนียมธนาคาร (Online Payment: Thai QR, Online Credit Card เท่านั้น)"),
    ("Students",        "จำนวนนักเรียนทั้งหมด (unique) ที่มี Invoice ในช่วงที่ filter"),
    ("Transactions",    "จำนวน Transaction ทั้งหมด"),
    ("Success Rate",    "อัตรา Transaction สำเร็จ (%)"),
]
for i, (k, v) in enumerate(stats):
    bg = LBLUE if i % 2 == 0 else WHITE
    add_data_row(t, [k, v], bg=bg, bold_first=True)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# Section 3 — Invoice นับเป็นรายได้
# ══════════════════════════════════════════════════════════════════════════════
add_section_title(doc, "3. Invoice ที่นับเป็นรายได้")

t = doc.add_table(rows=1, cols=2)
t.style = "Table Grid"
header_row(t, ["หัวข้อ", "รายละเอียด"])
set_col_widths(t, [5, 10.5])

invoice_rules = [
    ("Status ที่นับ",              "paid, sent, overdue, approved"),
    ("Status ที่ไม่นับ",           "draft, cancelled, rejected, wait"),
    ("ปีปัจจุบัน (Current Year)",  "= ปีการศึกษาล่าสุดที่มีข้อมูลใน Invoice (ตรวจจากข้อมูลจริงในระบบ)"),
]
for i, (k, v) in enumerate(invoice_rules):
    bg = LBLUE if i % 2 == 0 else WHITE
    add_data_row(t, [k, v], bg=bg, bold_first=True)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# Section 4 — Navigation Tabs
# ══════════════════════════════════════════════════════════════════════════════
add_section_title(doc, "4. Navigation Tabs")

t = doc.add_table(rows=1, cols=2)
t.style = "Table Grid"
header_row(t, ["Tab", "คำอธิบาย"])
set_col_widths(t, [4, 11.5])

tabs_list = [
    ("Tab 1 — Revenue Comparison",       "เปรียบเทียบรายได้ YoY / ToT (ดูรายละเอียด Section 5)"),
    ("Tab 2 — AVG Amount",               "รายได้เฉลี่ยต่อนักเรียน (ดูรายละเอียด Section 6–7)"),
    ("Tab 3 — No. of Transactions",      "จำนวน Transaction แยก Payment Method"),
    ("Tab 4 — Declined vs Successful",   "สถานะ Transaction (ระบบนี้ Online Payment ไม่มี Decline)"),
    ("Tab 5 — Bank Fees",                "ค่าธรรมเนียมธนาคาร (Thai QR + Online Credit Card เท่านั้น)"),
    ("Tab 6 — Net vs Gross Revenue",     "Waterfall แสดงผลกระทบส่วนลดต่อรายได้"),
]
for i, (k, v) in enumerate(tabs_list):
    bg = LBLUE if i % 2 == 0 else WHITE
    add_data_row(t, [k, v], bg=bg, bold_first=True)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# Section 5 — Filter Logic (General)
# ══════════════════════════════════════════════════════════════════════════════
add_section_title(doc, "5. เงื่อนไข Filter — กฎหลัก (ใช้กับทุก Tab: 1, 2, 3, 5)")

t = doc.add_table(rows=1, cols=3)
t.style = "Table Grid"
header_row(t, ["Year Filter", "Term Filter", "ผลลัพธ์"])
set_col_widths(t, [3.5, 3.5, 8.5])

rules = [
    ("All Years",    "All Terms",   "แสดงข้อมูลปีการศึกษาปัจจุบัน (ปีล่าสุดในระบบ)"),
    ("Filter Year",  "All Terms",   "แสดงปีที่เลือก + ย้อนหลัง 2 ปี (รวม 3 ปี)"),
    ("All Years",    "Filter Term", "แสดง Term นั้น + ย้อนหลัง 2 ปี (รวม 3 ปี)"),
    ("Filter Year",  "Filter Term", "แสดง Term นั้นใน 3 ปี โดยใช้ปีที่เลือกเป็นฐาน"),
]
for i, (y, t_val, res) in enumerate(rules):
    bg = LBLUE if i % 2 == 0 else WHITE
    add_data_row(t, [y, t_val, res], bg=bg, bold_first=True)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# Section 6 — Tab 1: Revenue Comparison Detail
# ══════════════════════════════════════════════════════════════════════════════
add_section_title(doc, "6. Tab 1 — Revenue Comparison: เงื่อนไขแสดงผล")

t = doc.add_table(rows=1, cols=3)
t.style = "Table Grid"
header_row(t, ["Filter", "Table / Chart", "Columns ที่แสดง"])
set_col_widths(t, [4, 4, 7.5])

t1_rules = [
    ("All Years + All Terms",  "Bar Chart",                  "Revenue ปีปัจจุบัน (1 set of bars per Year Group)"),
    ("Filter Year",            "Bar Chart",                  "Grouped bars 3 ปี เปรียบเทียบ (3 bars per Year Group)"),
    ("Filter Term",            "Bar Chart",                  "Revenue ปีปัจจุบัน Term นั้น (1 bar per Year Group)"),
    ("All Years + All Terms",  "Compare by Term",            "Terms ของปีปัจจุบัน เช่น 2025/26 T1, T2, T3"),
    ("Filter Year",            "Compare by Term",            "ทุก Term × 3 ปี เช่น 23/24 T1, 24/25 T1, 25/26 T1, T2, T3"),
    ("Filter Term",            "Compare by Term",            "Term นั้นใน 3 ปี เช่น 23/24 T1, 24/25 T1, 25/26 T1"),
    ("All Years + All Terms",  "Compare by Academic Year",   "ปีปัจจุบัน (1 column)"),
    ("Filter Year",            "Compare by Academic Year",   "3 ปี: ปีที่เลือก + 2 ปีก่อน"),
    ("Filter Term",            "Compare by Academic Year",   "ไม่แสดงข้อมูล (—)"),
]
for i, (f, tbl, cols) in enumerate(t1_rules):
    bg = LBLUE if i % 2 == 0 else WHITE
    add_data_row(t, [f, tbl, cols], bg=bg, bold_first=True)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# Section 7 — Tab 2: AVG Amount Filter Logic
# ══════════════════════════════════════════════════════════════════════════════
add_section_title(doc, "7. Tab 2 — AVG Amount: เงื่อนไขแสดงผล")

t = doc.add_table(rows=1, cols=3)
t.style = "Table Grid"
header_row(t, ["Filter", "Table", "Columns ที่แสดง"])
set_col_widths(t, [4, 5, 6.5])

t2_rules = [
    ("All Years + All Terms",  "Avg Revenue by Term",           "Terms ของปีปัจจุบัน"),
    ("Filter Year",            "Avg Revenue by Term",           "ทุก Term × 3 ปี"),
    ("Filter Term",            "Avg Revenue by Term",           "Term นั้นใน 3 ปี"),
    ("All Years + All Terms",  "Avg Revenue by Academic Year",  "ปีปัจจุบัน (1 column)"),
    ("Filter Year",            "Avg Revenue by Academic Year",  "3 ปี"),
    ("Filter Term",            "Avg Revenue by Academic Year",  "Term นั้น × 3 ปี"),
]
for i, (f, tbl, cols) in enumerate(t2_rules):
    bg = LBLUE if i % 2 == 0 else WHITE
    add_data_row(t, [f, tbl, cols], bg=bg, bold_first=True)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# Section 8 — AVG Amount: โครงสร้างตาราง
# ══════════════════════════════════════════════════════════════════════════════
add_section_title(doc, "8. Tab 2 — AVG Amount: โครงสร้างตาราง")

t = doc.add_table(rows=1, cols=3)
t.style = "Table Grid"
header_row(t, ["Column", "ชื่อ", "คำอธิบาย"])
set_col_widths(t, [3.5, 4, 8])

cols_detail = [
    ("Fixed",         "Year Group",    "ชื่อกลุ่มชั้นปี เช่น Pre-Nursery, Year 1 ... Year 13"),
    ("Per Term/Year", "Students",      "จำนวนนักเรียน (unique) ที่มี Invoice ใน Term/Year นั้น"),
    ("Per Term/Year", "Total amount",  "รายได้รวมทั้งหมดของ Year Group นั้นใน Period นั้น"),
    ("Per Term/Year", "Avg amount",    "Total amount ÷ Students = รายได้เฉลี่ยต่อนักเรียน 1 คน"),
]
for i, (col_type, name, desc) in enumerate(cols_detail):
    bg = LBLUE if i % 2 == 0 else WHITE
    add_data_row(t, [col_type, name, desc], bg=bg, bold_first=True)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# Section 9 — AVG Amount: แถว Total
# ══════════════════════════════════════════════════════════════════════════════
add_section_title(doc, "9. Tab 2 — AVG Amount: แถว Total (ต่อคน)")

t = doc.add_table(rows=1, cols=2)
t.style = "Table Grid"
header_row(t, ["Column", "วิธีคำนวณ"])
set_col_widths(t, [4, 11.5])

totals = [
    ("Students",      "ผลรวม unique students ใน Term/Year นั้น (ทุก Year Group)"),
    ("Total amount",  "ผลรวม Total amount ทุก Year Group"),
    ("Avg amount",    "Total amount รวม ÷ Students รวม"),
]
for i, (name, desc) in enumerate(totals):
    bg = LBLUE if i % 2 == 0 else WHITE
    add_data_row(t, [name, desc], bg=bg, bold_first=True)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# Section 10 — ตัวอย่างการคำนวณ
# ══════════════════════════════════════════════════════════════════════════════
add_section_title(doc, "10. Tab 2 — ตัวอย่างการคำนวณ (Year 7 / 2025/2026 Term 1)")

t = doc.add_table(rows=1, cols=3)
t.style = "Table Grid"
header_row(t, ["รายการ", "ค่า", "หมายเหตุ"])
set_col_widths(t, [4, 3, 8.5])

example = [
    ("Students",      "15 คน",    ""),
    ("Total amount",  "฿825,000", "รายได้รวมทุก Invoice ของ Year 7 ใน Term 1 ปี 2025/26"),
    ("Avg amount",    "฿55,000",  "825,000 ÷ 15 = 55,000 บาทต่อคน"),
]
for i, (name, val, note) in enumerate(example):
    bg = YELLOW if i == 2 else (LBLUE if i % 2 == 0 else WHITE)
    row = t.add_row()
    set_cell_bg(row.cells[0], bg); set_cell_text(row.cells[0], name, bold=True)
    set_cell_bg(row.cells[1], bg); set_cell_text(row.cells[1], val, align=WD_ALIGN_PARAGRAPH.RIGHT)
    set_cell_bg(row.cells[2], bg); set_cell_text(row.cells[2], note)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# Section 11 — Tab 3: No. of Transactions
# ══════════════════════════════════════════════════════════════════════════════
add_section_title(doc, "11. Tab 3 — No. of Transactions")

t = doc.add_table(rows=1, cols=2)
t.style = "Table Grid"
header_row(t, ["หัวข้อ", "รายละเอียด"])
set_col_widths(t, [5, 10.5])

tab3 = [
    ("ข้อมูลที่แสดง",   "จำนวน Transaction แยกตาม Payment Method"),
    ("ตาราง 1",         "Year Group × Payment Method Matrix — นับจำนวน Transaction"),
    ("ตาราง 2",         "Payment Method × Academic Year — เปรียบเทียบ YoY (ใช้ target years เดียวกับ Tab 1)"),
    ("Payment Methods", "Bank Transfer, Onsite Credit Card, Thai QR, Online Credit Card, Cheque, Bill Payment"),
    ("Filter Year",     "ตาราง YoY แสดง 3 ปี: ปีที่เลือก + 2 ปีก่อน"),
    ("All Years",       "ตาราง YoY แสดงแค่ปีปัจจุบัน (ปีล่าสุดในระบบ)"),
    ("Filter Term",     "กรองเฉพาะ Term นั้น × 3 ปี (เปรียบเทียบ YoY)"),
    ("Fee Type Filter", "ไม่มีผลต่อ Tab นี้ (payment records ไม่มี category field)"),
]
for i, (k, v) in enumerate(tab3):
    bg = LBLUE if i % 2 == 0 else WHITE
    add_data_row(t, [k, v], bg=bg, bold_first=True)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# Section 12 — Tab 4: Declined vs Successful
# ══════════════════════════════════════════════════════════════════════════════
add_section_title(doc, "12. Tab 4 — Declined vs Successful")

t = doc.add_table(rows=1, cols=2)
t.style = "Table Grid"
header_row(t, ["หัวข้อ", "รายละเอียด"])
set_col_widths(t, [5, 10.5])

tab4 = [
    ("หมายเหตุ",       "ระบบนี้เป็น Offline Payment ทั้งหมด → ไม่มี Declined"),
    ("ข้อมูลที่แสดง",  "จำนวน Successful Transaction แยกตาม Payment Method"),
    ("Success Rate",   "= 100% เสมอ (ไม่มี Gateway Decline)"),
]
for i, (k, v) in enumerate(tab4):
    bg = LBLUE if i % 2 == 0 else WHITE
    add_data_row(t, [k, v], bg=bg, bold_first=True)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# Section 13 — Tab 5: Bank Fees
# ══════════════════════════════════════════════════════════════════════════════
add_section_title(doc, "13. Tab 5 — Bank Fees")

t = doc.add_table(rows=1, cols=2)
t.style = "Table Grid"
header_row(t, ["หัวข้อ", "รายละเอียด"])
set_col_widths(t, [5, 10.5])

tab5 = [
    ("Online Payment ที่มีค่าธรรมเนียม", "Thai QR, Online Credit Card เท่านั้น"),
    ("ไม่นับ",                           "Bank Transfer, Cheque, Bill Payment, Onsite Credit Card"),
    ("ตาราง 1",                           "Bank × Term Matrix — ค่าธรรมเนียมแยกตาม Term (ใช้ target years เดียวกับ Tab 1)"),
    ("ตาราง 2",                           "Bank × Academic Year — เปรียบเทียบ YoY (Filter Term → ไม่แสดง)"),
    ("All Years + All Terms",             "ตาราง YoY แสดงแค่ปีปัจจุบัน"),
    ("Filter Year + All Terms",           "ตาราง YoY แสดง 3 ปี: ปีที่เลือก + 2 ปีก่อน"),
    ("Filter Term",                       "ตาราง Term แสดง Term นั้น × 3 ปี / ตาราง YoY ไม่แสดง (—)"),
]
for i, (k, v) in enumerate(tab5):
    bg = LBLUE if i % 2 == 0 else WHITE
    add_data_row(t, [k, v], bg=bg, bold_first=True)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# Section 14 — Tab 6: Net vs Gross Revenue
# ══════════════════════════════════════════════════════════════════════════════
add_section_title(doc, "14. Tab 6 — Net vs Gross Revenue")

t = doc.add_table(rows=1, cols=2)
t.style = "Table Grid"
header_row(t, ["หัวข้อ", "รายละเอียด"])
set_col_widths(t, [5, 10.5])

tab6 = [
    ("Fee Type Filter",   "ล็อกที่ 'All Fee Types' — ไม่สามารถเปลี่ยนได้บน Tab นี้ (เน้นแค่ค่าเทอม)"),
    ("Revenue Waterfall", "ไม่มี (ลบออกแล้ว)"),
    ("Gross Revenue",     "รายได้ก่อนหักส่วนลดทุกประเภท (= Subtotal)"),
    ("Net Revenue",       "= Gross Revenue − ผลรวม Deductions"),
]
for i, (k, v) in enumerate(tab6):
    bg = LBLUE if i % 2 == 0 else WHITE
    add_data_row(t, [k, v], bg=bg, bold_first=True)

doc.add_paragraph()

add_section_title(doc, "14a. Tab 6 — Revenue Breakdown by Year Group: โครงสร้างตาราง")

t = doc.add_table(rows=1, cols=2)
t.style = "Table Grid"
header_row(t, ["Column", "คำอธิบาย"])
set_col_widths(t, [5, 10.5])

breakdown_cols = [
    ("Year Group",                         "ชื่อกลุ่มชั้นปี เช่น Pre-Nursery, Year 1 ... Year 13"),
    ("Students",                           "จำนวนนักเรียน (unique) ที่มี Invoice"),
    ("Gross Revenue",                      "รายได้รวมก่อนหักส่วนลด"),
    ("[Discount columns จาก Discount Mgmt]", "แต่ละคอลัมน์แสดง ชื่อกลุ่มส่วนลด + อัตรา (% หรือ ฿) ใต้ชื่อ\nเช่น: Scholarships (20%), 2nd Child (5%), Staff Children (฿5,000)"),
    ("Net Revenue",                        "= Gross Revenue − รวมส่วนลดทุกประเภท"),
]
for i, (k, v) in enumerate(breakdown_cols):
    bg = LBLUE if i % 2 == 0 else WHITE
    add_data_row(t, [k, v], bg=bg, bold_first=True)

# ── Save ──────────────────────────────────────────────────────────────────────
out = "/Users/passkornnabangchang/Desktop/warp/kingcollenge 2/Kingcollegebackoffice-main/docs/Analytics_Dashboard_Spec.docx"
doc.save(out)
print(f"Saved: {out}")
